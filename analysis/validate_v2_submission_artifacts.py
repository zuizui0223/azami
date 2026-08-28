#!/usr/bin/env python3
"""Validate frozen v2 science and, when requested, the committed submission bundle."""
from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "submission"
RESULTS = ROOT / "analysis_outputs" / "v2_full27_environment_atlas_2026-08-27"
FIGURES = ROOT / "manuscript" / "figures" / "v2_submission"
TABLES = ROOT / "manuscript" / "tables" / "v2_submission"
TEXT_SUFFIXES = {".csv", ".json", ".md", ".py", ".toml", ".txt"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--science-only",
        action="store_true",
        help="validate frozen scientific reports without requiring the committed presentation bundle hashes to match edited manuscript sources",
    )
    return parser.parse_args()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def package_bytes(path: Path, name: str) -> bytes:
    if Path(name).suffix.lower() in TEXT_SUFFIXES:
        text = path.read_text(encoding="utf-8").replace("\r\n", "\n").replace("\r", "\n")
        return text.encode("utf-8")
    return path.read_bytes()


def source_path(bundle_name: str) -> Path:
    if bundle_name == "README.md":
        return SUBMISSION / "README.md"
    if bundle_name.startswith("figures/"):
        return FIGURES / bundle_name.removeprefix("figures/")
    if bundle_name.startswith("figure_source/"):
        return FIGURES / "source" / bundle_name.removeprefix("figure_source/")
    if bundle_name.startswith("tables/"):
        return TABLES / bundle_name.removeprefix("tables/")
    if bundle_name.startswith("data/"):
        return RESULTS / bundle_name.removeprefix("data/")
    if bundle_name.startswith("source/"):
        return ROOT / bundle_name.removeprefix("source/")
    if bundle_name == "code/pyproject.toml":
        return ROOT / "pyproject.toml"
    if bundle_name.startswith("code/"):
        return ROOT / "analysis" / bundle_name.removeprefix("code/")
    return SUBMISSION / bundle_name


def check(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def validate_frozen_science() -> int:
    validation_reports = [
        RESULTS / "v2_full27_environment_validation.json",
        RESULTS / "sampling" / "v2_full27_sampling_composition_validation.json",
        RESULTS / "v2_full27_sensitivities_validation.json",
    ]
    checks = 0
    for path in validation_reports:
        report = json.loads(path.read_text(encoding="utf-8"))
        check(report["status"] == "PASS", f"frozen validation is not PASS: {path.name}")
        check(report["n_checks"] == report["n_passed"], f"failed frozen check: {path.name}")
        checks += int(report["n_checks"])
    return checks


def validate_figure_table_contract() -> int:
    figure_report = json.loads((FIGURES / "figure_build_report.json").read_text(encoding="utf-8"))
    check(figure_report["status"] == "ok", "figure build report is not ok")
    main_stems = figure_report["main_figure_stems"]
    supporting_stems = figure_report["supporting_figure_stems"]
    check(len(main_stems) == 5, "Main figure count is not five")
    check(len(supporting_stems) == 4, "Supporting figure count is not four")
    for stem in main_stems + supporting_stems:
        check((FIGURES / f"{stem}.png").is_file(), f"missing PNG: {stem}")
        check((FIGURES / f"{stem}.pdf").is_file(), f"missing PDF: {stem}")

    table_report = json.loads((TABLES / "table_build_report.json").read_text(encoding="utf-8"))
    check(table_report["status"] == "PASS", "table build report is not PASS")
    expected_rows = {
        "Table_1_v2_analytical_roadmap.csv": 8,
        "Table_2_v2_candidate_evidence_chain.csv": 2,
        "Table_S1_v2_complete_27_endpoint_contract.csv": 27,
        "Table_S2_v2_cohort_geographic_filtering.csv": 5,
        "Table_S3_v2_complete_variance_decomposition.csv": 27,
        "Table_S4_v2_complete_within_taxon_environment_atlas.csv": 234,
        "Table_S5_v2_complete_among_taxon_environment_atlas.csv": 468,
        "Table_S6_v2_complete_cross_scale_classification.csv": 234,
        "Table_S7_v2_sampling_composition_robustness.csv": 134,
        "Table_S8_v2_spatial_historical_sequential_gate.csv": 36,
        "Table_S9_v2_secondary_whole_capitulum_summary.csv": 2,
        "Table_S10_v2_full_analytical_roadmap.csv": 8,
    }
    check(table_report["table_rows"] == expected_rows, "submission table row counts drifted")
    check((SUBMISSION / "Azami_Chapter1_v2_Supplementary_Tables.xlsx").is_file(), "supplementary workbook missing")
    audit_path = ROOT / "manuscript" / "GEB_V2_FIGURE_TABLE_AUDIT_2026-08-28.md"
    mapping_path = TABLES / "Manuscript_claim_figure_table_mapping.csv"
    check(audit_path.is_file(), "figure/table audit report missing")
    check(mapping_path.is_file(), "claim-to-figure/table mapping missing")
    mapping_text = mapping_path.read_text(encoding="utf-8")
    check("Detector-positive source geography" in mapping_text and "figure missing" in mapping_text, "unmapped detector-positive geography gap was hidden")

    active_paths = [
        *sorted((ROOT / "manuscript").glob("0[0-7]_*.md")),
        ROOT / "manuscript" / "SUBMISSION_MANUSCRIPT.md",
        ROOT / "manuscript" / "supplement" / "SUPPLEMENTARY_INFORMATION.md",
        ROOT / "analysis" / "build_v2_submission_figures.py",
        ROOT / "analysis" / "build_v2_submission_tables.py",
        ROOT / "analysis" / "build_v2_submission_package.py",
    ]
    forbidden = (
        "9 primary endpoints", "6,626", "3,725", "36 primary CHELSA tests",
        "SPDE-INLA", "32.9%", "23.2%", "69.3%",
    )
    for path in active_paths:
        text = path.read_text(encoding="utf-8")
        check(not any(token.lower() in text.lower() for token in forbidden), f"v1-only token in active submission source: {path}")
    supplement = (ROOT / "manuscript" / "supplement" / "SUPPLEMENTARY_INFORMATION.md").read_text(encoding="utf-8")
    check("52 audited" in supplement and "50 randomized" in supplement, "52-tree/50-randomized distinction missing")
    figure_builder = (ROOT / "analysis" / "build_v2_submission_figures.py").read_text(encoding="utf-8")
    check("fig.suptitle(" not in figure_builder, "figure-level title remains embedded in artwork")
    check('set_title("Figure ' not in figure_builder, "numbered figure title remains embedded in artwork")
    documents = {
        filename: Document(SUBMISSION / filename)
        for filename in ("Azami_Chapter1_v2_Main.docx", "Azami_Chapter1_v2_Supplement.docx")
    }
    for filename, document in documents.items():
        check(all(section.page_width < section.page_height for section in document.sections), f"non-portrait page remains in {filename}")

    main_document = documents["Azami_Chapter1_v2_Main.docx"]
    main_text = "\n".join(paragraph.text for paragraph in main_document.paragraphs)
    check("EAzami" not in main_text and "Azami Chapter 1" not in main_text, "internal project labels remain in the public-facing main manuscript")
    numbered_level_1 = {"Introduction", "Methods", "Results", "Discussion", "Conclusion and declarations"}
    numbered_level_2_parents = {"Methods", "Results", "Discussion"}
    current_level_1 = ""
    numbering_ids: set[str] = set()
    numbered_heading_count = 0
    unnumbered_heading_count = 0
    for paragraph in main_document.paragraphs:
        if paragraph.style is None or not paragraph.style.name.startswith("Heading "):
            continue
        style_level = int(paragraph.style.name.rsplit(" ", 1)[1])
        if style_level == 1:
            current_level_1 = paragraph.text
        should_be_numbered = (
            (style_level == 1 and paragraph.text in numbered_level_1)
            or (style_level == 2 and current_level_1 in numbered_level_2_parents)
        )
        properties = paragraph._p.pPr
        numbering = properties.find(qn("w:numPr")) if properties is not None else None
        if should_be_numbered:
            check(numbering is not None, f"body heading is not numbered: {paragraph.text}")
            level = numbering.find(qn("w:ilvl"))
            number_id = numbering.find(qn("w:numId"))
            check(level is not None and level.get(qn("w:val")) == str(style_level - 1), f"wrong heading level number: {paragraph.text}")
            check(number_id is not None, f"heading numbering id is absent: {paragraph.text}")
            numbering_ids.add(number_id.get(qn("w:val")))
            numbered_heading_count += 1
        else:
            check(numbering is None, f"front/back-matter heading must remain unnumbered: {paragraph.text}")
            unnumbered_heading_count += 1
    check(numbered_heading_count == 25, "numbered body-heading count drifted")
    check(len(numbering_ids) == 1, "body headings do not share one multilevel sequence")
    check(unnumbered_heading_count >= 13, "front/back-matter heading audit is incomplete")
    for section in main_document.sections:
        line_numbers = section._sectPr.find(qn("w:lnNumType"))
        check(line_numbers is not None, "main manuscript line numbering is absent")
        check(line_numbers.get(qn("w:countBy")) == "1", "main manuscript line numbering is not every line")
        check(line_numbers.get(qn("w:restart")) == "continuous", "main manuscript line numbering is not continuous")
    return 5 + 4 + len(expected_rows) + len(active_paths) + 16


def validate_committed_bundle() -> tuple[int, str]:
    manifest_path = SUBMISSION / "SUBMISSION_MANIFEST.json"
    sums_path = SUBMISSION / "SHA256SUMS.txt"
    zip_path = SUBMISSION / "Azami_Chapter1_v2_submission_candidate.zip"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    rows = manifest["files"]
    check(manifest["status"] == "scientific_hold_pending_independent_validation", "HOLD status changed")
    check(manifest["source_lane"] == "full27_full_environment_only", "source lane changed")
    check(len(rows) == len({row["path"] for row in rows}), "duplicate manifest paths")
    row_names = {str(row["path"]) for row in rows}
    check("source/manuscript/GEB_V2_FIGURE_TABLE_AUDIT_2026-08-28.md" in row_names, "audit report missing from bundle")
    check("tables/Manuscript_claim_figure_table_mapping.csv" in row_names, "claim mapping missing from bundle")

    expected_hashes: dict[str, str] = {}
    for row in rows:
        name = str(row["path"])
        path = source_path(name)
        check(path.is_file(), f"missing manifest source: {name}")
        data = package_bytes(path, name)
        check(len(data) == int(row["bytes"]), f"byte count drift: {name}")
        digest = sha256_bytes(data)
        check(digest == row["sha256"], f"SHA-256 drift: {name}")
        expected_hashes[name] = digest

    with ZipFile(zip_path) as archive:
        names = set(archive.namelist())
        expected_names = set(expected_hashes) | {"SUBMISSION_MANIFEST.json", "SHA256SUMS.txt"}
        check(names == expected_names, "ZIP inventory differs from manifest")
        for name, digest in expected_hashes.items():
            check(sha256_bytes(archive.read(name)) == digest, f"ZIP payload drift: {name}")
        check(
            archive.read("SUBMISSION_MANIFEST.json")
            == package_bytes(manifest_path, "SUBMISSION_MANIFEST.json"),
            "ZIP manifest differs from external manifest",
        )

    sums = {
        line.split("  ", 1)[1]: line.split("  ", 1)[0]
        for line in sums_path.read_text(encoding="utf-8").splitlines()
        if "  " in line
    }
    check(all(sums.get(name) == digest for name, digest in expected_hashes.items()), "SHA256SUMS payload mismatch")
    check(sums.get(zip_path.name) == sha256_bytes(zip_path.read_bytes()), "ZIP SHA-256 mismatch")

    forbidden = (
        "continuous_trait_reanalysis_v1",
        "reviewer_bias_control_v1",
        "pr69_",
        "legacy_lability",
        "trait_lability",
    )
    check(
        not any(any(token in name.lower() for token in forbidden) for name in expected_hashes),
        "legacy path leaked into bundle",
    )
    return len(rows), sha256_bytes(zip_path.read_bytes())


def main() -> None:
    args = parse_args()
    checks = validate_frozen_science()
    presentation_checks = validate_figure_table_contract()
    if args.science_only:
        print(json.dumps({
            "status": "PASS",
            "mode": "science_only",
            "frozen_validation_checks": checks,
            "figure_table_contract_checks": presentation_checks,
            "scientific_status": "frozen_v2_reports_pass",
        }, indent=2))
        return

    manifest_files, zip_digest = validate_committed_bundle()
    manifest = json.loads((SUBMISSION / "SUBMISSION_MANIFEST.json").read_text(encoding="utf-8"))
    print(json.dumps({
        "status": "PASS",
        "mode": "full_bundle",
        "manifest_files": manifest_files,
        "frozen_validation_checks": checks,
        "figure_table_contract_checks": presentation_checks,
        "zip_sha256": zip_digest,
        "scientific_status": manifest["status"],
    }, indent=2))


if __name__ == "__main__":
    main()
