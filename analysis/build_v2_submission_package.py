#!/usr/bin/env python3
"""Build blinded DOCX manuscripts and a hashed v2 submission-candidate bundle."""
from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "submission"
FIGURES = ROOT / "manuscript" / "figures" / "v2_submission"
TABLES = ROOT / "manuscript" / "tables" / "v2_submission"
RESULTS = ROOT / "analysis_outputs" / "v2_full27_environment_atlas_2026-08-27"
WORKBOOK = ROOT / "submission" / "Azami_Chapter1_v2_Supplementary_Tables.xlsx"

MAIN_SOURCES = [
    ROOT / "manuscript" / "00_title_abstract.md",
    ROOT / "manuscript" / "01_introduction.md",
    ROOT / "manuscript" / "02_methods.md",
    ROOT / "manuscript" / "03_results.md",
    ROOT / "manuscript" / "04_discussion.md",
    ROOT / "manuscript" / "05_conclusion_and_declarations.md",
    ROOT / "manuscript" / "06_references.md",
    ROOT / "manuscript" / "07_data_code_availability.md",
]

SUPPLEMENT_SOURCE = ROOT / "manuscript" / "supplement" / "SUPPLEMENTARY_INFORMATION.md"
TITLE = "Global continuous phenomics reveals trait- and scale-specific environmental geography of thistle capitula"

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--finalize", action="store_true", help="require PDFs and build final ZIP/manifest")
    return parser.parse_args()


def set_run_font(run, name: str = "Times New Roman", size: float = 11.5) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            label = token[1:token.index("]")]
            run = paragraph.add_run(label)
            run.font.color.rgb = None
        set_run_font(run, run.font.name or "Times New Roman", run.font.size.pt if run.font.size else 11.5)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]))


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    for style_name, size, bold in (
        ("Normal", 11.5, False),
        ("Heading 1", 15, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11.5, True),
    ):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal = document.styles["Normal"].paragraph_format
    normal.line_spacing = 1.35
    normal.space_after = Pt(5)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = "Blinded submission candidate; scientific HOLD retained."


def remove_heading_numbering(document: Document) -> None:
    """Remove inherited style numbering before applying manuscript-specific numbering."""
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        properties = document.styles[style_name]._element.get_or_add_pPr()
        numbering = properties.find(qn("w:numPr"))
        if numbering is not None:
            properties.remove(numbering)


def create_heading_numbering(document: Document) -> int:
    """Create genuine Word multilevel numbering for GEB-style body headings."""
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract.append(multi_level)

    for level, label in enumerate(("%1", "%1.%2", "%1.%2.%3")):
        level_element = OxmlElement("w:lvl")
        level_element.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level_element.append(start)
        number_format = OxmlElement("w:numFmt")
        number_format.set(qn("w:val"), "decimal")
        level_element.append(number_format)
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), label)
        level_element.append(level_text)
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "space")
        level_element.append(suffix)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        level_element.append(justification)
        abstract.append(level_element)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    num.append(abstract_reference)
    numbering.append(num)
    return num_id


def apply_heading_numbering(paragraph, num_id: int, level: int) -> None:
    """Attach a body heading to the shared multilevel numbering definition."""
    properties = paragraph._p.get_or_add_pPr()
    existing = properties.find(qn("w:numPr"))
    if existing is not None:
        properties.remove(existing)
    numbering = OxmlElement("w:numPr")
    indent_level = OxmlElement("w:ilvl")
    indent_level.set(qn("w:val"), str(level - 1))
    numbering_id = OxmlElement("w:numId")
    numbering_id.set(qn("w:val"), str(num_id))
    numbering.extend([indent_level, numbering_id])
    properties.append(numbering)


def enable_continuous_line_numbering(document: Document) -> None:
    """Add continuous reviewer line numbers to every main-manuscript section."""
    for section in document.sections:
        properties = section._sectPr
        existing = properties.find(qn("w:lnNumType"))
        if existing is not None:
            properties.remove(existing)
        line_numbers = OxmlElement("w:lnNumType")
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:distance"), "360")
        line_numbers.set(qn("w:restart"), "continuous")
        properties.append(line_numbers)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_markdown(
    document: Document,
    text: str,
    *,
    skip_title_scaffold: bool = False,
    heading_numbering_id: int | None = None,
    numbered_heading_levels: set[int] | None = None,
) -> None:
    lines = text.splitlines()
    paragraph_buffer: list[str] = []

    def flush() -> None:
        if paragraph_buffer:
            paragraph = document.add_paragraph()
            add_inline(paragraph, " ".join(value.strip() for value in paragraph_buffer))
            paragraph_buffer.clear()

    skipped_preferred_title = False
    for raw in lines:
        line = raw.rstrip()
        if skip_title_scaffold and line in {"# Title and abstract", "## Preferred title"}:
            continue
        if skip_title_scaffold and not skipped_preferred_title and line.startswith("**Global continuous phenomics"):
            skipped_preferred_title = True
            continue
        if not line.strip():
            flush()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush()
            level = len(heading.group(1))
            title = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", heading.group(2))
            if skip_title_scaffold and title == "Abstract":
                level = 1
            elif skip_title_scaffold and level == 3:
                level = 2
            paragraph = document.add_heading(title, level=level)
            if heading_numbering_id is not None and level in (numbered_heading_levels or set()):
                apply_heading_numbering(paragraph, heading_numbering_id, level)
            continue
        if line.startswith("- "):
            flush()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, line[2:].strip())
            continue
        if re.match(r"^\d+\.\s+", line):
            flush()
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", line))
            continue
        if line.startswith("> "):
            flush()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.35)
            run = paragraph.add_run(line[2:])
            run.italic = True
            set_run_font(run)
            continue
        if line.startswith("|"):
            raise ValueError("Markdown table detected; convert it explicitly before DOCX build")
        paragraph_buffer.append(line)
    flush()


def set_image_alt_text(shape, description: str) -> None:
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description.split(".", 1)[0])


def add_figure(document: Document, path: Path, caption: str, alt_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(6.3))
    set_image_alt_text(shape, alt_text)
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_paragraph.paragraph_format.keep_together = True
    caption_run = caption_paragraph.add_run(caption)
    caption_run.bold = True
    set_run_font(caption_run, size=10)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def mark_header_row(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        properties.append(header)
    header.set(qn("w:val"), "true")


def add_csv_table(
    document: Document,
    path: Path,
    caption: str,
    *,
    widths: list[float] | None = None,
    font_size: float = 7.4,
) -> None:
    rows = list(csv.reader(path.read_text(encoding="utf-8").splitlines()))
    if not rows:
        raise ValueError(f"Empty table: {path}")
    heading = document.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(caption)
    run.bold = True
    set_run_font(run, size=10)
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = widths is None
    mark_header_row(table.rows[0])
    for column, value in enumerate(rows[0]):
        cell = table.rows[0].cells[column]
        cell.text = value.replace("_", " ")
        set_cell_shading(cell, "0F5B66")
        for cell_run in cell.paragraphs[0].runs:
            cell_run.bold = True
            cell_run.font.color.rgb = RGBColor(255, 255, 255)
            set_run_font(cell_run, size=font_size)
    for values in rows[1:]:
        cells = table.add_row().cells
        for column, value in enumerate(values):
            cells[column].text = value
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[column], "EAF3F5")
            for cell_run in cells[column].paragraphs[0].runs:
                set_run_font(cell_run, size=font_size)
    if widths is not None:
        if len(widths) != len(rows[0]):
            raise ValueError(f"Column-width count does not match {path.name}")
        for row in table.rows:
            for column, width in enumerate(widths):
                row.cells[column].width = Inches(width)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 0
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0


def build_main(output: Path) -> Path:
    document = Document()
    style_document(document)
    remove_heading_numbering(document)
    enable_continuous_line_numbering(document)
    heading_numbering_id = create_heading_numbering(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(TITLE)
    title_run.bold = True
    set_run_font(title_run, size=16)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Blinded submission candidate — scientific HOLD retained")
    subrun.italic = True
    set_run_font(subrun, size=10)
    document.add_page_break()
    numbered_levels = {
        "01_introduction.md": {1},
        "02_methods.md": {1, 2},
        "03_results.md": {1, 2},
        "04_discussion.md": {1, 2},
        "05_conclusion_and_declarations.md": {1},
    }
    for index, source in enumerate(MAIN_SOURCES):
        add_markdown(
            document,
            source.read_text(encoding="utf-8"),
            skip_title_scaffold=index == 0,
            heading_numbering_id=heading_numbering_id,
            numbered_heading_levels=numbered_levels.get(source.name, set()),
        )
    document.add_page_break()
    document.add_heading("Tables", level=1)
    add_csv_table(
        document,
        TABLES / "Table_1_v2_analytical_roadmap.csv",
        "Table 1. Compact analytical roadmap for the global public-image study of thistle capitula. Each stage identifies its frozen scope, operation or decision gate, and linked display; the full eight-column roadmap is Table S10.",
        widths=[1.10, 1.35, 2.75, 1.00],
        font_size=7.5,
    )
    document.add_page_break()
    add_csv_table(
        document,
        TABLES / "Table_2_v2_candidate_evidence_chain.csv",
        "Table 2. Frozen evidence chain for the two among-taxon candidate patterns in the global thistle-capitulum cohort. Beta denotes a standardized coefficient; BH denotes Benjamini-Hochberg; P values are descriptive outputs of the predeclared gates. Stability under these controls is not evidence of adaptation, convergence or mechanism.",
        widths=[1.30, 1.05, 1.25, 1.35, 1.25],
        font_size=7.3,
    )
    document.add_page_break()
    document.add_heading("Figures", level=1)
    figures = [
        ("Figure_1_v2_measurement_pipeline.png", "Figure 1. From public photographs to deterministic continuous measurements in the global Carduus-Cirsium thistle-capitulum study. Three open-licensed photographs and their detector/crop examples are presentation/source provenance only; all numerical result statements are current v2.", "Real thistle photographs, detector bounding boxes, tight and context crops, colour mask, outline geometry, orientation and representative continuous outputs, followed by the current 27 registered, 22 measured and five unexecuted contract."),
        ("Figure_2_v2_geographic_sampling_domain.png", "Figure 2. Realized global sampling domain of the frozen public-image cohort of thistle capitula. (a) Disclosure-safe observation density, (b) locked filtering flow to 46,276 observations and 259 source-assigned taxa, and (c) the realized annual-temperature (BIO1) and annual-precipitation (BIO12) domain. The panels do not establish representativeness.", "World map of disclosure-safe observation density, locked filtering flow to 46,276 observations and 259 taxa, and the realized BIO1-BIO12 domain."),
        ("Figure_3_v2_taxon_mean_information_loss.png", "Figure 3. Visible image-phenotype variation compressed by source-assigned taxon means across the global public-image cohort of thistle capitula. Horizontal bars compare the raw observed sample with deterministic equal-replication resamples across 22 measured endpoints; the partition is not genetic variance.", "Horizontal bars compare raw and equal-replication variation below taxon means across 22 measured endpoints; the partition is not genetic variance."),
        ("Figure_4_v2_within_among_scale_atlas.png", "Figure 4. Endpoint-specific environmental alignment within and among source-assigned thistle taxa across the global public-image cohort. The complete 234-row heat map distinguishes both-scale, within-only, among-only, neither and not-comparable rows; joint hue is one circular inferential unit.", "Complete heat map of 234 both-scale, within-only, among-only, neither and not-comparable rows; hue is a joint circular unit and unexecuted rows remain not comparable."),
        ("Figure_5_v2_candidate_robustness.png", "Figure 5. Sequential attrition and robustness of the two final candidate patterns in the global thistle-capitulum cohort. (a) The 234-row grid is reduced through among-taxon multiplicity, sampling annotation, broad-space and 52-tree placement gates; (b-c) retain the coefficient and gate evidence for chroma-radiation and image-referenced orientation-precipitation. Stability is not evidence of adaptation, convergence or mechanism.", "Flow from 234 canonical rows through among-taxon multiplicity, sampling annotation, broad-space and 52-tree placement gates, plus the evidence retained for chroma-radiation and orientation-precipitation."),
    ]
    for filename, caption, alt in figures:
        add_figure(document, FIGURES / filename, caption, alt)
    add_page_number(document.sections[0].footer.paragraphs[0])
    destination = output / "Azami_Chapter1_v2_Main.docx"
    document.save(destination)
    return destination


def build_supplement(output: Path) -> Path:
    document = Document()
    style_document(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Supporting Information")
    run.bold = True
    set_run_font(run, size=16)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(TITLE)
    subtitle_run.italic = True
    set_run_font(subtitle_run, size=10)
    supplement_text = SUPPLEMENT_SOURCE.read_text(encoding="utf-8")
    # The source retains a reusable caption catalogue, but submission legends must
    # appear with their figures rather than on an earlier text-only page.
    supplement_body = supplement_text.split("## Supporting figure captions", 1)[0].rstrip()
    add_markdown(document, supplement_body)
    document.add_page_break()
    document.add_heading("Supplementary table workbook index", level=1)
    report = json.loads((TABLES / "table_build_report.json").read_text(encoding="utf-8"))
    index_rows = [("Table", "Rows", "Workbook sheet")]
    table_names = [
        ("Table S1", "Table_S1_v2_complete_27_endpoint_contract.csv", "S1 Endpoint contract"),
        ("Table S2", "Table_S2_v2_cohort_geographic_filtering.csv", "S2 Cohort filtering"),
        ("Table S3", "Table_S3_v2_complete_variance_decomposition.csv", "S3 Variance"),
        ("Table S4", "Table_S4_v2_complete_within_taxon_environment_atlas.csv", "S4 Within atlas"),
        ("Table S5", "Table_S5_v2_complete_among_taxon_environment_atlas.csv", "S5 Among atlas"),
        ("Table S6", "Table_S6_v2_complete_cross_scale_classification.csv", "S6 Cross-scale"),
        ("Table S7", "Table_S7_v2_sampling_composition_robustness.csv", "S7 Sampling"),
        ("Table S8", "Table_S8_v2_spatial_historical_sequential_gate.csv", "S8 Sequential gate"),
        ("Table S9", "Table_S9_v2_secondary_whole_capitulum_summary.csv", "S9 Secondary synthesis"),
        ("Table S10", "Table_S10_v2_full_analytical_roadmap.csv", "S10 Full roadmap"),
    ]
    for label, filename, sheet in table_names:
        index_rows.append((label, str(report["table_rows"][filename]), sheet))
    index = document.add_table(rows=1, cols=3)
    index.style = "Table Grid"
    mark_header_row(index.rows[0])
    for column, value in enumerate(index_rows[0]):
        index.rows[0].cells[column].text = value
        set_cell_shading(index.rows[0].cells[column], "0F5B66")
        for cell_run in index.rows[0].cells[column].paragraphs[0].runs:
            cell_run.bold = True
            cell_run.font.color.rgb = RGBColor(255, 255, 255)
            set_run_font(cell_run, size=9)
    for values in index_rows[1:]:
        cells = index.add_row().cells
        for column, value in enumerate(values):
            cells[column].text = value
            for cell_run in cells[column].paragraphs[0].runs:
                set_run_font(cell_run, size=9)
    document.add_page_break()
    document.add_heading("Supporting figures", level=1)
    add_figure(
        document,
        FIGURES / "Figure_S1_v2_endpoint_measurement_support.png",
        "Figure S1. Measurement support across the registered continuous-trait universe in the global public-image cohort of thistle capitula. Observation and source-assigned taxon coverage are shown for every measured endpoint; explicit unexecuted endpoints remain in the inventory. The display diagnoses data support and does not repeat the main taxon-mean information-loss result.",
        "Observation and taxon coverage for all measured endpoints; unexecuted rows remain explicit in the endpoint inventory.",
    )
    add_figure(
        document,
        FIGURES / "Figure_S2_v2_sampling_composition_audit.png",
        "Figure S2. Sampling-composition sensitivity of all globally supported rows entering the audit in the global public-image cohort of thistle capitula. Bars show the minimum retained effect-magnitude ratio across declared perturbations for each selected within- and among-taxon association; colour distinguishes direction-stable from direction-unstable rows. This is the full selected-row audit rather than only the two final candidates, and it contains no post-selection significance tests.",
        "Minimum retained effect-magnitude ratio for every selected within- and among-taxon row, with direction-instability flagged.",
    )
    add_figure(
        document,
        FIGURES / "Figure_S3_v2_spatial_diagnostic_surface.png",
        "Figure S3. Broad-space and residual-spatial diagnostic surface for all globally supported rows entering spatial sensitivity in the global public-image cohort of thistle capitula. The horizontal axis shows spatial-association support as -log10(permutation P), and the vertical axis shows residual eight-nearest-neighbour Moran P; point shape separates within- from among-taxon models and colour distinguishes full-gate passage from failure. The display exposes why many initially supported rows do not enter the final candidate set.",
        "Spatial permutation support versus residual Moran diagnostic for all rows entering the broad-space gate, with full-gate passage distinguished.",
    )
    add_figure(
        document,
        FIGURES / "Figure_S4_v2_historical_placement_stability.png",
        "Figure S4. Frozen 52-tree historical-placement sensitivity for the two among-taxon candidate patterns in the global thistle-capitulum cohort. Standardized Pagel-lambda phylogenetic generalized least-squares coefficients, P values and lambda estimates are shown for two deterministic and 50 randomized within-genus placements per candidate. The display is a placement sensitivity, not a resolved species-tree analysis.",
        "Three-panel strip plots show PGLS coefficients, P values and Pagel lambda across two deterministic and 50 randomized placement trees for chroma-radiation and orientation-precipitation.",
    )
    add_page_number(document.sections[0].footer.paragraphs[0])
    destination = output / "Azami_Chapter1_v2_Supplement.docx"
    document.save(destination)
    return destination


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


TEXT_SUFFIXES = {".csv", ".json", ".md", ".py", ".toml", ".txt"}


def package_bytes(path: Path, name: str) -> bytes:
    """Return cross-platform package bytes for tracked text and exact binary bytes."""
    if Path(name).suffix.lower() in TEXT_SUFFIXES:
        text = path.read_text(encoding="utf-8").replace("\r\n", "\n").replace("\r", "\n")
        return text.encode("utf-8")
    return path.read_bytes()


def bundle_sources(output: Path) -> list[tuple[Path, str]]:
    files: list[tuple[Path, str]] = []
    files.append((output / "README.md", "README.md"))
    for path in sorted(output.glob("Azami_Chapter1_v2_*.*")):
        if path.suffix.lower() in {".docx", ".pdf", ".xlsx"}:
            files.append((path, path.name))
    for path in sorted(FIGURES.glob("Figure_*.*")):
        if path.suffix.lower() in {".png", ".pdf"}:
            files.append((path, f"figures/{path.name}"))
    for path in sorted((FIGURES / "source").glob("*")):
        if path.is_file():
            files.append((path, f"figure_source/{path.name}"))
    for path in sorted(TABLES.glob("*")):
        if path.is_file():
            files.append((path, f"tables/{path.name}"))
    for path in sorted(RESULTS.rglob("*")):
        if path.is_file():
            files.append((path, f"data/{path.relative_to(RESULTS).as_posix()}"))
    for path in MAIN_SOURCES + [
        SUPPLEMENT_SOURCE,
        ROOT / "manuscript/final_claims.json",
        ROOT / "manuscript/EXTERNAL_COMPLETION_GATES.md",
        ROOT / "manuscript/GEB_V2_FIGURE_TABLE_AUDIT_2026-08-28.md",
    ]:
        files.append((path, f"source/{path.relative_to(ROOT).as_posix()}"))
    for path in (
        ROOT / "analysis/ch1/hypothesis_recovery_status.json",
        ROOT / "analysis/ch1/pipeline.json",
        ROOT / "analysis/ch1/submission_contract.json",
        ROOT / "analysis/ch1/v2_full27_environment_atlas_contract.json",
        ROOT / "analysis/ch1/v2_full27_sampling_composition_sensitivity_contract.json",
    ):
        files.append((path, f"source/{path.relative_to(ROOT).as_posix()}"))
    for path in (
        ROOT / "analysis/run_geb_v2_full27_environment_atlas.py",
        ROOT / "analysis/run_geb_v2_full27_sampling_composition_sensitivity.py",
        ROOT / "analysis/run_geb_v2_full27_spatial_sensitivity.py",
        ROOT / "analysis/run_geb_v2_full27_historical_sensitivity.py",
        ROOT / "analysis/validate_geb_v2_full27_environment_atlas.py",
        ROOT / "analysis/validate_geb_v2_full27_sampling_composition.py",
        ROOT / "analysis/validate_geb_v2_full27_sensitivities.py",
        ROOT / "analysis/validate_v2_submission_artifacts.py",
        ROOT / "analysis/validate_manuscript_citations.py",
        ROOT / "analysis/build_v2_submission_figures.py",
        ROOT / "analysis/build_v2_submission_figure_inputs.py",
        ROOT / "analysis/build_v2_submission_tables.py",
    ):
        files.append((path, f"code/{path.relative_to(ROOT / 'analysis').as_posix()}"))
    files.append((ROOT / "pyproject.toml", "code/pyproject.toml"))
    return files


def write_deterministic_zip(destination: Path, files: list[tuple[Path, str]]) -> None:
    with ZipFile(destination, "w", compression=ZIP_DEFLATED, compresslevel=9) as archive:
        for source, name in sorted(files, key=lambda item: item[1]):
            info = ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.external_attr = 0o644 << 16
            archive.writestr(info, package_bytes(source, name))


def finalize(output: Path) -> None:
    required = [
        output / "Azami_Chapter1_v2_Main.docx",
        output / "Azami_Chapter1_v2_Main.pdf",
        output / "Azami_Chapter1_v2_Supplement.docx",
        output / "Azami_Chapter1_v2_Supplement.pdf",
        output / "Azami_Chapter1_v2_Supplementary_Tables.xlsx",
    ]
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        raise SystemExit("Missing rendered submission files:\n- " + "\n- ".join(missing))
    files = bundle_sources(output)
    manifest = {
        "schema_version": 1,
        "release_label": "Azami Chapter 1 v2 submission candidate",
        "status": "scientific_hold_pending_independent_validation",
        "source_freeze_tag": "azami-ch1-v2-2026-08-27",
        "source_lane": "full27_full_environment_only",
        "claim_ceiling": "adaptive-pattern candidates under current sequential controls; not demonstrated adaptation or mechanism",
        "files": [
            {
                "path": name,
                "bytes": len(package_bytes(source, name)),
                "sha256": sha256_bytes(package_bytes(source, name)),
            }
            for source, name in sorted(files, key=lambda item: item[1])
        ],
    }
    manifest_path = output / "SUBMISSION_MANIFEST.json"
    manifest_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    sums_path = output / "SHA256SUMS.txt"
    sums_path.write_text(
        "\n".join(f'{row["sha256"]}  {row["path"]}' for row in manifest["files"]) + "\n",
        encoding="utf-8",
    )
    zip_path = output / "Azami_Chapter1_v2_submission_candidate.zip"
    zip_files = files + [
        (manifest_path, "SUBMISSION_MANIFEST.json"),
        (sums_path, "SHA256SUMS.txt"),
    ]
    write_deterministic_zip(zip_path, zip_files)
    with sums_path.open("a", encoding="utf-8") as handle:
        handle.write(f"{sha256(zip_path)}  {zip_path.name}\n")
    print(json.dumps({
        "status": "ok",
        "main_docx": str(required[0]),
        "supplement_docx": str(required[2]),
        "bundle": str(zip_path),
        "bundle_sha256": sha256(zip_path),
        "n_manifest_files": len(manifest["files"]),
    }, indent=2))


def main() -> None:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    for path in MAIN_SOURCES + [SUPPLEMENT_SOURCE]:
        text = path.read_text(encoding="utf-8")
        if re.search(r"\b(TODO|TBD|PLACEHOLDER)\b", text, flags=re.IGNORECASE):
            raise SystemExit(f"Placeholder token in {path}")
    if args.finalize:
        finalize(args.output_dir)
    else:
        workbook_destination = args.output_dir / WORKBOOK.name
        if WORKBOOK.resolve() != workbook_destination.resolve():
            shutil.copy2(WORKBOOK, workbook_destination)
        build_main(args.output_dir)
        build_supplement(args.output_dir)
        print(json.dumps({"status": "docx_built", "output": str(args.output_dir)}, indent=2))


if __name__ == "__main__":
    main()
